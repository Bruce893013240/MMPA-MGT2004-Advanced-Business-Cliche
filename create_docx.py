#!/usr/bin/env python3
"""
Convert MGT2004 Study Notes HTML to DOCX format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_highlight_box(doc, text):
    """Add a highlighted box with left border style"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)  # Dark blue
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Add a border/shading effect via paragraph style
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), '2F5496')
    pBdr.append(left)
    pPr.append(pBdr)
    return para

def add_key_point(doc, text):
    """Add a key point with emphasis"""
    para = doc.add_paragraph()
    run = para.add_run("★ " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    para.paragraph_format.left_indent = Inches(0.3)
    return para

def add_bullet_list(doc, items):
    """Add a bullet list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
    return para

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text

    return table

def create_docx():
    """Create the DOCX document from HTML content"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('MGT2004 高级战略管理概念', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Advanced Concepts in Strategic Management — 学习笔记')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph('多伦多大学 MMPA 项目 | Professor Kevin Yousie\n本笔记涵盖4个核心Sessions，保留重要英文术语以帮助考试准备。')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== Session 1 =====
    doc.add_heading('Session 1: Governance and the Role of the Board', 1)
    doc.add_paragraph('董事会角色与公司治理')

    doc.add_heading('Board of Directors: Role and Composition', 2)

    add_highlight_box(doc, 'Governance vs Management — 治理与管理\n'
                          '董事会的角色是治理(GOVERNANCE)而非管理(MANAGEMENT)。'
                          '董事会负责监督，而不直接参与日常运营。')

    doc.add_heading("The Board's Role — Directing the Affairs of the Company", 3)

    board_duties = [
        'Selection of the CEO — 选择并监督CEO，确保其胜任且道德地运营公司',
        'Ensure ethical operations — 确保管理层以有效且道德的方式运营，为股东创造价值',
        'Financial oversight — 监督财务报告的公正性，确保及时披露',
        'Independent audit — 聘请独立会计师事务所审计财务报表',
        'Shape governance — 塑造公司治理，推荐最佳董事候选人',
        'Compensation policies — 制定薪酬政策，确定高管薪酬，建立绩效补偿目标',
        'Shareholder concerns — 适当回应股东关切',
        'Fair treatment — 确保公司与员工、客户、供应商等公平相处'
    ]
    add_bullet_list(doc, board_duties)

    doc.add_heading("Director's Legal Duties", 2)
    doc.add_paragraph('董事的三大法律义务')

    add_highlight_box(doc, 'Duty of Care (注意义务)\n'
                          '董事必须以合理谨慎的人的标准履行职责，包括：熟悉公司业务、参加董事会会议、审查财务报表、提出明智的问题。')

    add_highlight_box(doc, 'Business Judgement Rule (商业判断规则)\n'
                          '只要董事出于善意、以合理信息为依据、且理性地认为决策符合公司最佳利益，'
                          '即使决策最终失败，也免于承担责任。')

    add_highlight_box(doc, 'Duty of Loyalty (忠诚义务)\n'
                          '董事必须将公司利益置于个人利益之上，避免利益冲突，不得利用职位谋取私利。')

    doc.add_heading('Shareholder Activism at Canadian Pacific', 2)
    doc.add_paragraph('案例：加拿大太平洋铁路股东维权')
    doc.add_paragraph('背景：Pershing Square Capital Management持有约14%股份，试图获得公司控制权，要求更换CEO和多名董事。')

    doc.add_heading('Lessons Learned', 3)

    lessons = [
        '即使业绩良好，董事也必须保持警惕，确保为股东创造最优结果',
        '董事会必须有应对特殊情况的能力（时间、精力投入）',
        '董事会构成至关重要——需具备应对特殊情况的知识和经验，可借助外部专家',
        '股东维权活动日益普遍，不会消失',
        '董事必须履行注意义务和忠诚义务，否则可能面临诉讼',
        '机构投资者有能力影响投票结果（短期利益、代理投票）',
        '股东激进主义者倾向于短期视角，可能以牺牲长期投资为代价提升股价',
        '有效的领导力（CEO/管理层和董事会）可以推动卓越成果',
        '有效的沟通在特殊情况下非常重要'
    ]
    add_bullet_list(doc, lessons)

    doc.add_page_break()

    # ===== Session 2 =====
    doc.add_heading('Session 2: Strategic Management', 1)
    doc.add_paragraph('战略管理')

    doc.add_heading('Strategy-Making, Strategy-Executing Process (5 Stages)', 2)
    doc.add_paragraph('战略制定与执行流程（5个阶段）')

    stages = [
        'Stage 1: 选择公司将要开展的业务和想要追求的愿景 — 选择公司将要开展的业务和想要追求的愿景',
        'Stage 2: 制定长期目标 — 制定长期目标',
        'Stage 3: 制定战略 — 制定实现目标的战略',
        'Stage 4: 执行战略 — 执行选定的战略',
        'Stage 5: 评估结果 — 评估结果并必要时进行修正'
    ]
    add_bullet_list(doc, stages)

    doc.add_heading('Vision and Mission Statements', 2)
    doc.add_paragraph('愿景与使命陈述')

    add_highlight_box(doc, 'Vision (愿景)\n'
                          '公司未来想要成为什么样的描述。它是关于"我们想成为什么"的问题。')

    add_highlight_box(doc, 'Mission (使命)\n'
                          '公司的业务定义——公司做什么和为谁做。它是关于"我们现在是谁"的问题。')

    doc.add_heading('Core Values and Strategic Intent', 2)

    add_highlight_box(doc, 'Core Values (核心价值观)\n'
                          '指导和塑造公司行为的持久信念和原则，如诚信、创新、客户至上等。')

    add_highlight_box(doc, 'Strategic Intent (战略意图)\n'
                          '对公司竞争地位的雄心勃勃但可实现的抱负。需要包含具体的衡量指标。')

    doc.add_heading('Sustainability and Strategy', 2)

    add_highlight_box(doc, 'Sustainability (可持续性)\n'
                          '公司满足当前需求同时不损害未来世代满足其需求的能力的承诺。'
                          '包括环境可持续性、社会责任和公司治理（ESG）三个维度。')

    doc.add_page_break()

    # ===== Session 3 =====
    doc.add_heading('Session 3: Competitive Strategy', 1)
    doc.add_paragraph('竞争战略')

    doc.add_heading('Building Competitive Advantage', 2)

    add_highlight_box(doc, 'Competitive Advantage (竞争优势)\n'
                          '公司相对于竞争对手的独特能力，使其能够为客户创造更多价值或以更低成本创造同等价值。')

    add_highlight_box(doc, 'Strategic Groups (战略群体)\n'
                          '同一行业内采用相似战略或具有相似特征的公司群体。'
                          '同一战略群体内的公司比不同群体内的公司更直接地相互竞争。')

    doc.add_heading('Offensive vs Defensive Strategies', 2)
    doc.add_paragraph('进攻型 vs 防御型战略')

    doc.add_paragraph('进攻型战略 (Offensive Strategy): 挑战者公司用来从领先者那里夺取市场份额的行动。')
    offensive = [
        'Direct Attack — 直接进攻：正面竞争',
        'Frontal Attack — 侧翼进攻：在竞争对手的弱点发起攻击',
        'Encirclement — 包围战略：从多个方向同时进攻',
        'Bypass Attack — 迂回战略：避开直接竞争，在不同领域建立优势',
        'Guerrilla Warfare — 游击战：小规模、间歇性的攻击'
    ]
    add_bullet_list(doc, offensive)

    doc.add_paragraph('防御型战略 (Defensive Strategy): 市场领导者用来保护市场份额和盈利能力的行动。')
    defensive = [
        'Fortification — 筑垒防御：加强现有地位',
        'Counterattack — 反击：主动出击对抗挑战者',
        'Mobility Barrier — 移动壁垒：限制挑战者的选择',
        'Contract — 收缩战略：放弃弱势领域，聚焦核心优势'
    ]
    add_bullet_list(doc, defensive)

    doc.add_heading('Blue Ocean Strategy', 2)

    add_highlight_box(doc, 'Blue Ocean Strategy (蓝海战略)\n'
                          '通过创造新的、无竞争的市场空间来规避竞争。与"红海"（充满竞争的市场）相对。')

    add_key_point(doc, '蓝海战略通过价值创新来创造新需求，打破价值-成本权衡')

    doc.add_heading('First-Mover vs Late-Mover Advantages', 2)

    add_highlight_box(doc, 'First-Mover Advantages (先动优势)\n'
                          '最早进入市场的公司获得的优势，包括：品牌忠诚度、客户转换成本、网络效应、规模经济等。')

    add_highlight_box(doc, 'Late-Mover Advantages (后动优势)\n'
                          '后进入市场的公司获得的优势，包括：学习先动者的经验、避免早期错误、观察市场需求等。')

    doc.add_heading('Mergers and Acquisitions (M&A)', 2)

    add_highlight_box(doc, 'Horizontal Integration/Mergers (横向整合/并购)\n'
                          '在同一业务层面、同一行业中的公司之间的并购。可以快速获得市场份额和能力，但面临整合风险。')

    add_highlight_box(doc, 'Vertical Integration (纵向整合)\n'
                          '公司收购或整合其供应商（后向整合）或分销商/客户（前向整合）。'
                          '后向整合（Backward Integration）：收购供应商\n'
                          '前向整合（Forward Integration）：收购分销商/客户')

    add_highlight_box(doc, 'Outsourcing (外包)\n'
                          '将非核心活动外包给第三方，以提高效率、降低成本或获取专业能力。')

    doc.add_heading('Strategic Alliances and Partnerships', 2)

    add_highlight_box(doc, 'Strategic Alliance (战略联盟)\n'
                          '两个或多个公司之间的合作协议，通常涉及共享资源、能力或知识，以实现共同目标。')

    doc.add_page_break()

    # ===== Session 4 =====
    doc.add_heading('Session 4: International Strategy', 1)
    doc.add_paragraph('国际战略')

    doc.add_heading("Porter's Diamond of National Advantage", 2)

    add_highlight_box(doc, "Porter's Diamond of National Advantage (波特国家竞争优势钻石模型)\n"
                          '解释为什么某些国家在特定行业具有竞争优势的框架，包括四个关键因素：')

    diamond_factors = [
        'Factor Conditions (要素条件) — 一个国家在特定竞争中所需的生产要素（如熟练劳动力、基础设施）',
        'Demand Conditions (需求条件) — 本国市场的需求特征（规模、成熟度、挑剔度）',
        'Related and Supporting Industries (相关和支持产业) — 供应商和相关产业的竞争力',
        'Firm Strategy, Structure, and Rivalry (企业战略、结构和竞争) — 企业的组织方式、管理实践和竞争强度'
    ]
    add_bullet_list(doc, diamond_factors)

    doc.add_heading('Five Modes of International Entry', 2)
    doc.add_paragraph('五种国际进入模式')

    entry_modes = [
        'Exporting (出口) — 在本国生产产品，然后销往其他国家。风险最低，但面临贸易壁垒和物流挑战。',
        'Licensing (许可) — 允许外国公司使用你的知识产权（如专利、商标），收取许可费。风险低，但控制有限。',
        'Franchising (特许经营) — 类似许可，但包括完整的商业模式和服务标准。常见于餐饮和酒店业。',
        'Acquisition (收购) — 购买外国公司的全部或大部分股权，快速获得市场准入和能力。',
        'Greenfield (绿地投资) — 从零开始在外国建立新的业务。风险最高，但提供完全控制和灵活性。'
    ]
    add_bullet_list(doc, entry_modes)

    doc.add_heading('Three International Strategies', 2)

    strategies = [
        'Global Strategy (全球战略) — 标准化的产品/服务在全球范围内统一运作。成本效率高，但文化适应性低。',
        'Multidomestic Strategy (多国本土化战略) — 根据每个国家的市场条件调整产品/服务。文化适应性好，但成本高。',
        'Transnational Strategy (跨国战略) — 试图同时实现全球效率和本地适应性的平衡。最复杂但潜在收益最大。'
    ]
    add_bullet_list(doc, strategies)

    doc.add_heading('Foreign Exchange Risk', 2)
    doc.add_paragraph('外汇风险')

    add_highlight_box(doc, 'Transaction Risk (交易风险)\n'
                          '因汇率变动导致实际交易现金流价值波动的风险。例如，出口商收到外币付款时汇率不利。')

    add_highlight_box(doc, 'Translation Risk (折算风险)\n'
                          '将海外子公司财务报表折算为本币时，因汇率变动导致的报告收益波动。')

    add_highlight_box(doc, 'Economic Risk (经济风险)\n'
                          '汇率变动对 公司长期价值和竞争地位的战略性影响。')

    doc.add_page_break()

    # ===== Exam Preparation =====
    doc.add_heading('Exam Preparation Section', 1)
    doc.add_paragraph('考试准备部分')

    doc.add_heading('Exam Information', 2)

    exam_info = [
        '考试时长：55分钟',
        '考试形式：选择题（约30-35道）',
        '考试内容：覆盖4个Sessions的所有内容',
        '注意事项：需提前15分钟入场，迟到者将被拒绝入场',
        '考试工具：可使用非编程型计算器（Non-programmable Calculator）'
    ]
    add_bullet_list(doc, exam_info)

    doc.add_heading('Key Study Areas', 2)

    study_areas = [
        'Board of Directors: 董事会职责、董事法律义务（三个Duty）',
        'Strategic Management: 战略制定与执行流程5阶段、愿景/使命/核心价值观',
        'Competitive Strategy: 竞争优势、进攻vs防御战略、蓝海vs红海战略',
        'International Strategy: 波特钻石模型、五种进入模式、三种国际战略、外汇风险'
    ]
    add_bullet_list(doc, study_areas)

    doc.add_page_break()

    # ===== Glossary =====
    doc.add_heading('Glossary — 专业词汇表', 1)

    headers = ['English Term', '中文解释']
    glossary = [
        ('Governance', '公司治理'),
        ('Board of Directors', '董事会'),
        ('Duty of Care', '注意义务'),
        ('Business Judgement Rule', '商业判断规则'),
        ('Duty of Loyalty', '忠诚义务'),
        ('Strategic Intent', '战略意图'),
        ('Core Values', '核心价值观'),
        ('Competitive Advantage', '竞争优势'),
        ('Strategic Groups', '战略群体'),
        ('Blue Ocean Strategy', '蓝海战略'),
        ('First-Mover Advantage', '先动优势'),
        ('Horizontal Integration', '横向整合'),
        ('Vertical Integration', '纵向整合'),
        ('Outsourcing', '外包'),
        ('Strategic Alliance', '战略联盟'),
        ("Porter's Diamond", '波特钻石模型'),
        ('Exporting', '出口'),
        ('Licensing', '许可'),
        ('Franchising', '特许经营'),
        ('Acquisition', '收购'),
        ('Greenfield Investment', '绿地投资'),
        ('Global Strategy', '全球战略'),
        ('Multidomestic Strategy', '多国本土化战略'),
        ('Transnational Strategy', '跨国战略'),
        ('Transaction Risk', '交易风险'),
        ('Translation Risk', '折算风险'),
        ('Economic Risk', '经济风险')
    ]

    add_table(doc, headers, glossary)

    # Save the document
    doc.save('/workspace/MGT2004_Study_Notes.docx')
    print("DOCX file created successfully: /workspace/MGT2004_Study_Notes.docx")

if __name__ == '__main__':
    create_docx()